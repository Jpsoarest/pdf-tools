from fastapi import FastAPI, File, UploadFile, HTTPException, Form
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pypdf import PdfReader, PdfWriter
import pikepdf
import os
import sys
from pathlib import Path
import uuid
import zipfile
import shutil
from typing import Optional, List, Any
import json
import io
import traceback
from PIL import Image
import math

# ─── OCR-HTR-LOCAL Integration ────────────────────────────────────────────────
OCR_LOCAL_PATH = Path("C:/Users/joaop/Downloads/Projetos/OCR-HTR-LOCAL")
if OCR_LOCAL_PATH.exists():
    sys.path.insert(0, str(OCR_LOCAL_PATH))

app = FastAPI(title="PDF Tools API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "https://4caxiastools.com.br",
        "https://www.4caxiastools.com.br",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=[
        "X-Original-Size", "X-Compressed-Size", "X-Reduction-Percent",
        "X-Total-Pages", "X-Files-Generated", "X-Pages-Converted",
        "X-Reordered-Pages", "X-Images-Converted", "X-Output-Pages",
        "X-Pages-Processed", "X-Tables-Found", "X-Rotated-Pages",
        "X-Extracted-Pages", "X-Removed-Pages", "X-Input-Size",
        "X-Output-Size", "X-Operations-Applied", "X-Fields-Added",
        "X-Replaced-Text"
    ]
)

TEMP_DIR = Path("temp")
TEMP_DIR.mkdir(exist_ok=True)

MAX_FILE_SIZE = 200 * 1024 * 1024


def cleanup_temp(*paths: Path):
    for p in paths:
        try:
            if p.is_file():
                os.remove(p)
            elif p.is_dir():
                shutil.rmtree(p)
        except Exception:
            pass


def validate_pdf(filename: str):
    if not filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Arquivo precisa ser PDF")


def parse_page_range(pages_str: str, total: int) -> List[int]:
    pages_str = pages_str.strip()
    if not pages_str or pages_str.lower() == 'all':
        return list(range(1, total + 1))

    result = []
    for part in pages_str.split(','):
        part = part.strip()
        if '-' in part:
            s, e = map(int, part.split('-'))
            for n in range(s, min(e + 1, total + 1)):
                result.append(n)
        else:
            n = int(part)
            if 1 <= n <= total:
                result.append(n)
    return sorted(set(result))


def parse_pdf_color(value: Any, default=(0, 0, 0)) -> tuple[float, float, float]:
    if isinstance(value, str):
        color = value.strip().lstrip("#")
        if len(color) == 3:
            color = "".join(ch * 2 for ch in color)
        if len(color) == 6:
            try:
                return tuple(int(color[i:i + 2], 16) / 255 for i in (0, 2, 4))
            except ValueError:
                return default

    if isinstance(value, list) and len(value) >= 3:
        try:
            return tuple(max(0, min(1, float(value[i]))) for i in range(3))
        except (TypeError, ValueError):
            return default

    return default


def parse_bool(value: Any) -> bool:
    if isinstance(value, bool):
        return value
    if isinstance(value, str):
        return value.strip().lower() in ("1", "true", "yes", "sim", "on")
    return bool(value)


def safe_field_name(value: Any, prefix: str, index: int, used: set[str]) -> str:
    raw = str(value or f"{prefix}_{index + 1}").strip()
    safe = "".join(ch if ch.isalnum() or ch in ("_", "-", ".") else "_" for ch in raw)
    safe = (safe[:80] or f"{prefix}_{index + 1}")
    base = safe
    suffix = 2
    while safe in used:
        safe = f"{base}_{suffix}"
        suffix += 1
    used.add(safe)
    return safe


def operation_rect(fitz_module, op: dict, page_rect, min_width=8, min_height=8):
    try:
        x = float(op.get("x", 0))
        y = float(op.get("y", 0))
        width = max(min_width, float(op.get("width", min_width)))
        height = max(min_height, float(op.get("height", min_height)))
    except (TypeError, ValueError):
        raise HTTPException(status_code=400, detail="Operacao contem coordenadas invalidas")

    x0 = max(0, min(page_rect.width - min_width, x))
    y0 = max(0, min(page_rect.height - min_height, y))
    x1 = max(x0 + min_width, min(page_rect.width, x0 + width))
    y1 = max(y0 + min_height, min(page_rect.height, y0 + height))
    return fitz_module.Rect(x0, y0, x1, y1)


def pdf_color_to_hex(value: Any) -> str:
    if isinstance(value, int):
        return f"#{(value >> 16) & 255:02X}{(value >> 8) & 255:02X}{value & 255:02X}"
    return "#111827"


def save_upload(file: UploadFile) -> tuple[Path, bytes]:
    content = file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="Arquivo excede 200MB")
    input_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}"
    os.makedirs(input_path.parent, exist_ok=True)
    with open(input_path, "wb") as f:
        f.write(content)
    return input_path, content


@app.get("/")
def read_root():
    return {"status": "ok", "message": "PDF Tools API esta rodando"}


# ─── COMPRIMIR PDF ────────────────────────────────────────────────────────────

@app.post("/compress-pdf")
async def compress_pdf(file: UploadFile = File(...)):
    validate_pdf(file.filename)
    input_id = str(uuid.uuid4())
    output_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}.pdf"
    output_path = TEMP_DIR / f"{output_id}_compressed.pdf"

    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)

        with pikepdf.Pdf.open(input_path) as pdf:
            pdf.save(output_path, compress_streams=True,
                     object_stream_mode=pikepdf.ObjectStreamMode.generate)

        original_size = os.path.getsize(input_path)
        compressed_size = os.path.getsize(output_path)

        if compressed_size >= original_size:
            os.remove(output_path)
            return FileResponse(
                path=input_path,
                media_type="application/pdf",
                filename=f"compressed_{file.filename}",
                headers={
                    "X-Original-Size": str(original_size),
                    "X-Compressed-Size": str(original_size),
                    "X-Reduction-Percent": "0.0"
                }
            )

        reduction = ((original_size - compressed_size) / original_size) * 100
        os.remove(input_path)

        return FileResponse(
            path=output_path,
            media_type="application/pdf",
            filename=f"compressed_{file.filename}",
            headers={
                "X-Original-Size": str(original_size),
                "X-Compressed-Size": str(compressed_size),
                "X-Reduction-Percent": f"{reduction:.1f}"
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        cleanup_temp(input_path, output_path)
        raise HTTPException(status_code=500, detail=f"Erro ao processar PDF: {str(e)}")


# ─── MESCLAR PDF ──────────────────────────────────────────────────────────────

@app.post("/merge-pdf")
async def merge_pdf(files: list[UploadFile] = File(...)):
    if len(files) < 2:
        raise HTTPException(status_code=400, detail="Envie pelo menos 2 arquivos PDF")
    for f in files:
        validate_pdf(f.filename)

    output_id = str(uuid.uuid4())
    output_path = TEMP_DIR / f"{output_id}_merged.pdf"
    temp_files = []

    try:
        for file in files:
            input_id = str(uuid.uuid4())
            input_path = TEMP_DIR / f"{input_id}.pdf"
            temp_files.append(input_path)
            content = await file.read()
            with open(input_path, "wb") as fw:
                fw.write(content)

        merger = PdfWriter()
        for pdf_path in temp_files:
            reader = PdfReader(pdf_path)
            for page in reader.pages:
                merger.add_page(page)

        with open(output_path, "wb") as of:
            merger.write(of)

        total_pages = len(merger.pages)
        cleanup_temp(*temp_files)

        return FileResponse(
            path=output_path,
            media_type="application/pdf",
            filename="merged.pdf",
            headers={"X-Total-Pages": str(total_pages)}
        )
    except Exception as e:
        cleanup_temp(*temp_files, output_path)
        raise HTTPException(status_code=500, detail=f"Erro ao mesclar PDFs: {str(e)}")


# ─── DIVIDIR PDF ──────────────────────────────────────────────────────────────

@app.post("/merge-pdf-batch")
async def merge_pdf_batch(files: list[UploadFile] = File(...), manifest: str = Form(...)):
    if not files:
        raise HTTPException(status_code=400, detail="Envie os arquivos PDF dos lotes")

    try:
        groups = json.loads(manifest)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Manifesto de lotes invalido")

    if not isinstance(groups, list) or not groups:
        raise HTTPException(status_code=400, detail="Adicione pelo menos um lote pronto")

    for file in files:
        validate_pdf(file.filename)

    input_paths = []
    output_paths = []
    zip_path = TEMP_DIR / f"{uuid.uuid4()}_merged_batch.zip"

    try:
        for file in files:
            input_path = TEMP_DIR / f"{uuid.uuid4()}.pdf"
            content = await file.read()
            with open(input_path, "wb") as output:
                output.write(content)
            input_paths.append(input_path)

        seen_names = set()
        zip_entries = []
        for position, group in enumerate(groups, start=1):
            if not isinstance(group, dict):
                raise HTTPException(status_code=400, detail="Lote invalido")

            raw_name = str(group.get("filename", "")).strip()
            indexes = group.get("indexes")
            if not raw_name or not isinstance(indexes, list) or len(indexes) < 2:
                raise HTTPException(status_code=400, detail="Cada lote precisa de nome e ao menos 2 PDFs")

            safe_name = "".join(
                ch if ch.isalnum() or ch in (" ", "-", "_", ".") else "_"
                for ch in raw_name
            ).strip(" .") or f"mesclagem_{position}"
            if not safe_name.lower().endswith(".pdf"):
                safe_name += ".pdf"
            if safe_name.lower() in seen_names:
                raise HTTPException(status_code=400, detail=f"Nome duplicado no lote: {safe_name}")
            seen_names.add(safe_name.lower())

            selected_paths = []
            for index in indexes:
                if not isinstance(index, int) or index < 0 or index >= len(input_paths):
                    raise HTTPException(status_code=400, detail="Referencia de arquivo invalida no lote")
                selected_paths.append(input_paths[index])

            output_path = TEMP_DIR / f"{uuid.uuid4()}_merged.pdf"
            writer = PdfWriter()
            for pdf_path in selected_paths:
                reader = PdfReader(pdf_path)
                for page in reader.pages:
                    writer.add_page(page)
            with open(output_path, "wb") as output:
                writer.write(output)
            output_paths.append(output_path)
            zip_entries.append((output_path, safe_name))

        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as archive:
            for output_path, output_name in zip_entries:
                archive.write(output_path, output_name)

        cleanup_temp(*input_paths, *output_paths)
        return FileResponse(
            path=zip_path,
            media_type="application/zip",
            filename="mesclagens_prontas.zip",
            headers={"X-Files-Generated": str(len(zip_entries))}
        )
    except HTTPException:
        cleanup_temp(*input_paths, *output_paths, zip_path)
        raise
    except Exception as e:
        cleanup_temp(*input_paths, *output_paths, zip_path)
        raise HTTPException(status_code=500, detail=f"Erro ao processar lote de mesclagens: {str(e)}")


@app.post("/split-pdf")
async def split_pdf(
    file: UploadFile = File(...),
    mode: str = "all",
    ranges: str = ""
):
    validate_pdf(file.filename)
    input_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}.pdf"
    output_folder = TEMP_DIR / f"{input_id}_split"
    output_folder.mkdir(exist_ok=True)

    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)

        reader = PdfReader(input_path)
        total_pages = len(reader.pages)
        output_files = []

        if mode == "all":
            for i in range(total_pages):
                writer = PdfWriter()
                writer.add_page(reader.pages[i])
                out = output_folder / f"page_{i + 1}.pdf"
                with open(out, "wb") as of:
                    writer.write(of)
                output_files.append(out)
        elif mode == "range":
            for range_str in ranges.split(','):
                writer = PdfWriter()
                range_str = range_str.strip()
                if '-' in range_str:
                    start, end = map(int, range_str.split('-'))
                    for pn in range(start - 1, min(end, total_pages)):
                        writer.add_page(reader.pages[pn])
                    out = output_folder / f"pages_{start}-{end}.pdf"
                else:
                    pn = int(range_str) - 1
                    if 0 <= pn < total_pages:
                        writer.add_page(reader.pages[pn])
                    out = output_folder / f"page_{range_str}.pdf"
                with open(out, "wb") as of:
                    writer.write(of)
                output_files.append(out)
        elif mode == "specific":
            for pn in [int(p.strip()) for p in ranges.split(',')]:
                if 1 <= pn <= total_pages:
                    writer = PdfWriter()
                    writer.add_page(reader.pages[pn - 1])
                    out = output_folder / f"page_{pn}.pdf"
                    with open(out, "wb") as of:
                        writer.write(of)
                    output_files.append(out)

        zip_id = str(uuid.uuid4())
        zip_path = TEMP_DIR / f"{zip_id}_split.zip"
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            for pf in output_files:
                zf.write(pf, pf.name)

        cleanup_temp(input_path, output_folder, *output_files)

        return FileResponse(
            path=zip_path,
            media_type="application/zip",
            filename=f"split_{file.filename.replace('.pdf', '')}.zip",
            headers={
                "X-Total-Pages": str(total_pages),
                "X-Files-Generated": str(len(output_files))
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        cleanup_temp(input_path, output_folder)
        raise HTTPException(status_code=500, detail=f"Erro ao dividir: {str(e)}")


# ─── PDF PARA JPG ─────────────────────────────────────────────────────────────

@app.post("/pdf-to-jpg")
async def pdf_to_jpg(
    file: UploadFile = File(...),
    quality: str = Form("high"),
    page_mode: str = Form("all"),
    pages: str = Form(""),
):
    validate_pdf(file.filename)
    try:
        import fitz
    except ImportError:
        raise HTTPException(status_code=500, detail="Dependencia nao instalada: pymupdf")

    dpi_map = {"low": 72, "medium": 150, "high": 300}
    dpi = dpi_map.get(quality, 150)
    zoom = dpi / 72

    input_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}.pdf"
    output_folder = TEMP_DIR / f"{input_id}_jpg"
    output_folder.mkdir(exist_ok=True)

    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)

        doc = fitz.open(str(input_path))
        total_pages = doc.page_count

        if page_mode == "specific" and pages.strip():
            page_indexes = parse_page_range(pages, total_pages)
            page_indexes = [p - 1 for p in page_indexes]
        else:
            page_indexes = list(range(total_pages))

        image_paths = []
        mat = fitz.Matrix(zoom, zoom)
        for idx in page_indexes:
            page = doc[idx]
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img_path = output_folder / f"pagina_{idx + 1:03d}.jpg"
            pix.save(str(img_path))
            image_paths.append(img_path)

        doc.close()

        if not image_paths:
            raise HTTPException(status_code=400, detail="Nenhuma pagina valida para converter")

        zip_id = str(uuid.uuid4())
        zip_path = TEMP_DIR / f"{zip_id}_jpg.zip"
        base_name = file.filename.replace('.pdf', '')
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zf:
            for ip in image_paths:
                zf.write(ip, ip.name)

        cleanup_temp(input_path, *image_paths, output_folder)

        return FileResponse(
            path=zip_path,
            media_type="application/zip",
            filename=f"{base_name}_imagens.zip",
            headers={"X-Pages-Converted": str(len(image_paths))}
        )
    except HTTPException:
        raise
    except Exception as e:
        cleanup_temp(input_path, output_folder)
        raise HTTPException(status_code=500, detail=f"Erro ao converter: {str(e)}")


# ─── PROTEGER PDF ─────────────────────────────────────────────────────────────

@app.post("/protect-pdf")
async def protect_pdf(
    file: UploadFile = File(...),
    password: str = Form(...)
):
    validate_pdf(file.filename)
    if len(password) < 4:
        raise HTTPException(status_code=400, detail="A senha deve ter pelo menos 4 caracteres")

    input_id = str(uuid.uuid4())
    output_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}.pdf"
    output_path = TEMP_DIR / f"{output_id}_protected.pdf"

    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)

        pdf = pikepdf.Pdf.open(input_path)
        pdf.save(output_path, encryption=pikepdf.Encryption(user=password, owner=password))
        pdf.close()

        original_size = os.path.getsize(input_path)
        protected_size = os.path.getsize(output_path)
        cleanup_temp(input_path)

        return FileResponse(
            path=output_path,
            media_type="application/pdf",
            filename=f"protected_{file.filename}",
            headers={
                "X-Original-Size": str(original_size),
                "X-Protected-Size": str(protected_size)
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        cleanup_temp(input_path, output_path)
        raise HTTPException(status_code=500, detail=f"Erro ao proteger PDF: {str(e)}")


# ─── REMOVER SENHA DO PDF ─────────────────────────────────────────────────────

@app.post("/unlock-pdf")
async def unlock_pdf(
    file: UploadFile = File(...),
    password: str = Form(...)
):
    validate_pdf(file.filename)
    input_id = str(uuid.uuid4())
    output_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}.pdf"
    output_path = TEMP_DIR / f"{output_id}_unlocked.pdf"

    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)

        try:
            pdf = pikepdf.Pdf.open(input_path, password=password)
        except pikepdf.PasswordError:
            raise HTTPException(status_code=401, detail="Senha incorreta")

        pdf.save(output_path)
        pdf.close()

        original_size = os.path.getsize(input_path)
        unlocked_size = os.path.getsize(output_path)
        cleanup_temp(input_path)

        return FileResponse(
            path=output_path,
            media_type="application/pdf",
            filename=f"unlocked_{file.filename}",
            headers={
                "X-Original-Size": str(original_size),
                "X-Unlocked-Size": str(unlocked_size)
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        cleanup_temp(input_path, output_path)
        raise HTTPException(status_code=500, detail=f"Erro ao remover senha: {str(e)}")


# ─── REORDENAR PDF ────────────────────────────────────────────────────────────

@app.post("/reorder-pdf")
async def reorder_pdf(
    file: UploadFile = File(...),
    order: str = Form(...)
):
    validate_pdf(file.filename)
    input_id = str(uuid.uuid4())
    output_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}.pdf"
    output_path = TEMP_DIR / f"{output_id}_reordered.pdf"

    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)

        reader = PdfReader(input_path)
        total_pages = len(reader.pages)

        try:
            new_order = json.loads(order)
        except json.JSONDecodeError:
            raise HTTPException(status_code=400, detail="Formato de ordem invalido. Use JSON array.")

        if len(new_order) != total_pages:
            raise HTTPException(status_code=400, detail=f"A ordem deve conter exatamente {total_pages} paginas")

        if sorted(new_order) != list(range(1, total_pages + 1)):
            raise HTTPException(status_code=400, detail="A ordem deve conter todas as paginas de 1 a N, sem duplicatas ou faltas")

        for p in new_order:
            if p < 1 or p > total_pages:
                raise HTTPException(status_code=400, detail=f"Pagina {p} fora do intervalo (1 a {total_pages})")

        writer = PdfWriter()
        for page_num in new_order:
            writer.add_page(reader.pages[page_num - 1])

        with open(output_path, "wb") as of:
            writer.write(of)

        cleanup_temp(input_path)

        return FileResponse(
            path=output_path,
            media_type="application/pdf",
            filename=f"reordered_{file.filename}",
            headers={
                "X-Total-Pages": str(total_pages),
                "X-Reordered-Pages": str(total_pages)
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        cleanup_temp(input_path, output_path)
        raise HTTPException(status_code=500, detail=f"Erro ao reordenar: {str(e)}")


# ─── IMAGEM PARA PDF ──────────────────────────────────────────────────────────

@app.post("/images-to-pdf")
async def images_to_pdf(
    files: List[UploadFile] = File(...),
    page_size: str = Form("auto"),
    orientation: str = Form("auto"),
    margin_mm: int = Form(10),
    fit: str = Form("contain"),
):
    if not files:
        raise HTTPException(status_code=400, detail="Envie pelo menos uma imagem")

    allowed_ext = {'.jpg', '.jpeg', '.png', '.webp'}
    for f in files:
        ext = os.path.splitext(f.filename)[1].lower()
        if ext not in allowed_ext:
            raise HTTPException(status_code=400, detail=f"Formato nao suportado: {f.filename}")

    output_id = str(uuid.uuid4())
    output_path = TEMP_DIR / f"{output_id}_images.pdf"
    temp_files = []

    try:
        pil_images = []
        for f in files:
            content = await f.read()
            img = Image.open(io.BytesIO(content))
            if img.mode in ('RGBA', 'LA'):
                background = Image.new('RGB', img.size, (255, 255, 255))
                background.paste(img, mask=img.split()[-1])
                img = background
            elif img.mode == 'P':
                background = Image.new('RGB', img.size, (255, 255, 255))
                img = img.convert('RGBA')
                background.paste(img, mask=img.split()[-1])
                img = background
            elif img.mode != 'RGB':
                img = img.convert('RGB')
            pil_images.append(img)

        page_widths = {
            'a4': (595, 842),
            'letter': (612, 792),
            'auto': None,
        }

        margin_pts = margin_mm * 2.834645
        pdf_images = []

        for img in pil_images:
            if page_size == 'auto':
                pw, ph = img.size
                if orientation == 'landscape' and pw < ph:
                    pw, ph = ph, pw
                elif orientation == 'portrait' and pw > ph:
                    pw, ph = ph, pw
                available_w = pw + 2 * margin_pts
                available_h = ph + 2 * margin_pts
            else:
                pw, ph = page_widths.get(page_size, (595, 842))
                if orientation == 'landscape':
                    pw, ph = ph, pw

                available_w = pw
                available_h = ph

            img_w, img_h = img.size
            draw_w = available_w - 2 * margin_pts
            draw_h = available_h - 2 * margin_pts

            if fit == 'contain':
                ratio = min(draw_w / img_w, draw_h / img_h)
            else:
                ratio = max(draw_w / img_w, draw_h / img_h)

            new_w = int(img_w * ratio)
            new_h = int(img_h * ratio)
            img = img.resize((new_w, new_h), Image.LANCZOS)

            pdf_img = Image.new('RGB', (int(available_w), int(available_h)), (255, 255, 255))
            offset_x = int((available_w - new_w) / 2)
            offset_y = int((available_h - new_h) / 2)
            pdf_img.paste(img, (offset_x, offset_y))
            pdf_images.append(pdf_img)

        if pdf_images:
            pdf_images[0].save(
                output_path,
                save_all=True,
                append_images=pdf_images[1:],
                format='PDF'
            )

        return FileResponse(
            path=output_path,
            media_type="application/pdf",
            filename="images.pdf",
            headers={
                "X-Images-Converted": str(len(files)),
                "X-Output-Pages": str(len(pdf_images))
            }
        )
    except Exception as e:
        traceback.print_exc()
        cleanup_temp(output_path, *temp_files)
        raise HTTPException(status_code=500, detail=f"Erro ao converter imagens: {str(e)}")


# ─── PDF PARA EXCEL ───────────────────────────────────────────────────────────

@app.post("/pdf-to-excel")
async def pdf_to_excel(
    file: UploadFile = File(...),
    mode: str = Form("tables"),
    pages: str = Form(""),
):
    validate_pdf(file.filename)
    input_id = str(uuid.uuid4())
    output_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}.pdf"
    output_path = TEMP_DIR / f"{output_id}.xlsx"

    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)

        try:
            import pdfplumber
            import pandas as pd
            from openpyxl import Workbook
        except ImportError:
            raise HTTPException(status_code=500, detail="Dependencias nao instaladas: pdfplumber, pandas, openpyxl")

        reader = PdfReader(input_path)
        total_pages = len(reader.pages)

        if pages.strip():
            page_list = parse_page_range(pages, total_pages)
        else:
            page_list = list(range(1, total_pages + 1))

        tables_found = 0

        with pdfplumber.open(input_path) as pdf:
            if mode == 'tables':
                with pd.ExcelWriter(output_path, engine='openpyxl') as writer_obj:
                    has_data = False
                    for page_num in page_list:
                        page = pdf.pages[page_num - 1]
                        tables = page.extract_tables()
                        if tables:
                            for table_idx, table in enumerate(tables):
                                if table and any(any(cell for cell in row) for row in table):
                                    headers = table[0] if table else []
                                    data = table[1:] if len(table) > 1 else []
                                    df = pd.DataFrame(data, columns=headers if headers and all(headers) else None)
                                    sheet_name = f"pag{page_num}_tab{table_idx + 1}"[:31]
                                    df.to_excel(writer_obj, sheet_name=sheet_name, index=False)
                                    tables_found += 1
                                    has_data = True
                        else:
                            text = page.extract_text()
                            if text:
                                lines = text.split('\n')
                                df = pd.DataFrame({'linha': range(1, len(lines) + 1), 'texto': lines})
                                sheet_name = f"pag{page_num}_texto"[:31]
                                df.to_excel(writer_obj, sheet_name=sheet_name, index=False)
                                has_data = True

                    if not has_data:
                        raise HTTPException(status_code=422, detail="Nenhuma tabela ou texto encontrado neste PDF")
            else:
                rows = []
                with pdfplumber.open(input_path) as pdf:
                    for page_num in page_list:
                        page = pdf.pages[page_num - 1]
                        text = page.extract_text()
                        if text:
                            for line_num, line in enumerate(text.split('\n'), 1):
                                rows.append({'pagina': page_num, 'linha': line_num, 'texto': line})

                if not rows:
                    raise HTTPException(status_code=422, detail="Nenhum texto encontrado neste PDF")

                df = pd.DataFrame(rows)
                df.to_excel(output_path, index=False)

        cleanup_temp(input_path)

        return FileResponse(
            path=output_path,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            filename=f"{file.filename.replace('.pdf', '')}.xlsx",
            headers={
                "X-Pages-Processed": str(len(page_list)),
                "X-Tables-Found": str(tables_found)
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        cleanup_temp(input_path, output_path)
        raise HTTPException(status_code=500, detail=f"Erro ao converter para Excel: {str(e)}")


# ─── GIRAR PDF ────────────────────────────────────────────────────────────────

@app.post("/rotate-pdf")
async def rotate_pdf(
    file: UploadFile = File(...),
    rotation: int = Form(...),
    pages: str = Form("all"),
):
    validate_pdf(file.filename)
    if rotation not in (90, 180, 270):
        raise HTTPException(status_code=400, detail="Rotacao invalida. Use 90, 180 ou 270")

    input_id = str(uuid.uuid4())
    output_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}.pdf"
    output_path = TEMP_DIR / f"{output_id}_rotated.pdf"

    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)

        reader = PdfReader(input_path)
        total_pages = len(reader.pages)

        page_list = parse_page_range(pages, total_pages)

        writer = PdfWriter()
        for i in range(total_pages):
            page = reader.pages[i]
            if (i + 1) in page_list:
                page.rotate(rotation)
            writer.add_page(page)

        with open(output_path, "wb") as of:
            writer.write(of)

        cleanup_temp(input_path)

        return FileResponse(
            path=output_path,
            media_type="application/pdf",
            filename=f"rotated_{file.filename}",
            headers={
                "X-Total-Pages": str(total_pages),
                "X-Rotated-Pages": str(len(page_list))
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        cleanup_temp(input_path, output_path)
        raise HTTPException(status_code=500, detail=f"Erro ao girar PDF: {str(e)}")


# ─── EXTRAIR PAGINAS ──────────────────────────────────────────────────────────

@app.post("/extract-pages")
async def extract_pages(
    file: UploadFile = File(...),
    pages: str = Form(...),
):
    validate_pdf(file.filename)
    input_id = str(uuid.uuid4())
    output_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}.pdf"
    output_path = TEMP_DIR / f"{output_id}_extracted.pdf"

    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)

        reader = PdfReader(input_path)
        total_pages = len(reader.pages)
        page_list = parse_page_range(pages, total_pages)

        if not page_list:
            raise HTTPException(status_code=400, detail="Nenhuma pagina valida informada")

        writer = PdfWriter()
        for pn in page_list:
            writer.add_page(reader.pages[pn - 1])

        with open(output_path, "wb") as of:
            writer.write(of)

        cleanup_temp(input_path)

        return FileResponse(
            path=output_path,
            media_type="application/pdf",
            filename=f"extracted_{file.filename}",
            headers={
                "X-Total-Pages": str(total_pages),
                "X-Extracted-Pages": str(len(page_list))
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        cleanup_temp(input_path, output_path)
        raise HTTPException(status_code=500, detail=f"Erro ao extrair paginas: {str(e)}")


# ─── REMOVER PAGINAS ──────────────────────────────────────────────────────────

@app.post("/remove-pages")
async def remove_pages(
    file: UploadFile = File(...),
    pages_to_remove: str = Form(...),
):
    validate_pdf(file.filename)
    input_id = str(uuid.uuid4())
    output_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}.pdf"
    output_path = TEMP_DIR / f"{output_id}_removed.pdf"

    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)

        reader = PdfReader(input_path)
        total_pages = len(reader.pages)
        remove_list = parse_page_range(pages_to_remove, total_pages)

        if len(remove_list) >= total_pages:
            raise HTTPException(status_code=400, detail="Nao e possivel remover todas as paginas")

        remove_set = set(remove_list)
        writer = PdfWriter()
        kept = 0
        for i in range(total_pages):
            if (i + 1) not in remove_set:
                writer.add_page(reader.pages[i])
                kept += 1

        if kept == 0:
            raise HTTPException(status_code=400, detail="Nao e possivel remover todas as paginas")

        with open(output_path, "wb") as of:
            writer.write(of)

        cleanup_temp(input_path)

        return FileResponse(
            path=output_path,
            media_type="application/pdf",
            filename=f"without_pages_{file.filename}",
            headers={
                "X-Total-Pages": str(total_pages),
                "X-Removed-Pages": str(len(remove_list))
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        cleanup_temp(input_path, output_path)
        raise HTTPException(status_code=500, detail=f"Erro ao remover paginas: {str(e)}")


# ─── COMPRIMIR IMAGEM ─────────────────────────────────────────────────────────

@app.post("/compress-image")
async def compress_image(
    file: UploadFile = File(...),
    quality: int = Form(80),
    target_format: str = Form("original"),
):
    allowed_ext = {'.jpg', '.jpeg', '.png', '.webp'}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed_ext:
        raise HTTPException(status_code=400, detail="Formato nao suportado. Envie JPG, PNG ou WebP")

    if quality < 10 or quality > 100:
        raise HTTPException(status_code=400, detail="Qualidade deve ser entre 10 e 100")

    input_id = str(uuid.uuid4())
    output_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}{ext}"
    output_ext = ext if target_format == "original" else f".{target_format}"
    output_path = TEMP_DIR / f"{output_id}_compressed{output_ext}"

    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)

        img = Image.open(input_path)
        original_size = len(content)

        if img.mode == 'RGBA' and output_ext in ('.jpg', '.jpeg'):
            background = Image.new('RGB', img.size, (255, 255, 255))
            background.paste(img, mask=img.split()[3])
            img = background
        elif img.mode == 'RGBA':
            pass
        elif img.mode != 'RGB' and output_ext in ('.jpg', '.jpeg'):
            img = img.convert('RGB')

        save_kwargs = {}
        if output_ext in ('.jpg', '.jpeg'):
            save_kwargs = {'format': 'JPEG', 'quality': quality, 'optimize': True}
        elif output_ext == '.png':
            save_kwargs = {'format': 'PNG', 'optimize': True}
        elif output_ext == '.webp':
            save_kwargs = {'format': 'WEBP', 'quality': quality}

        if img.mode == 'RGBA' and output_ext == '.png':
            save_kwargs['format'] = 'PNG'

        img.save(output_path, **save_kwargs)

        output_size = os.path.getsize(output_path)
        cleanup_temp(input_path)

        return FileResponse(
            path=output_path,
            media_type=f"image/{output_ext.replace('.', '')}",
            filename=f"compressed_{os.path.splitext(file.filename)[0]}{output_ext}",
            headers={
                "X-Original-Size": str(original_size),
                "X-Compressed-Size": str(output_size),
                "X-Reduction-Percent": f"{((original_size - output_size) / original_size * 100):.1f}" if original_size > 0 else "0"
            }
        )
    except Exception as e:
        cleanup_temp(input_path, output_path)
        raise HTTPException(status_code=500, detail=f"Erro ao comprimir imagem: {str(e)}")


# ─── CONVERTER IMAGEM ─────────────────────────────────────────────────────────

@app.post("/convert-image")
async def convert_image(
    file: UploadFile = File(...),
    target_format: str = Form(...),
):
    allowed_ext = {'.jpg', '.jpeg', '.png', '.webp'}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed_ext:
        raise HTTPException(status_code=400, detail="Formato nao suportado. Envie JPG, PNG ou WebP")

    if target_format not in ('jpg', 'png', 'webp'):
        raise HTTPException(status_code=400, detail="Formato de destino invalido. Use jpg, png ou webp")

    if ext == f'.{target_format}':
        raise HTTPException(status_code=400, detail="A imagem ja esta no formato desejado")

    input_id = str(uuid.uuid4())
    output_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}{ext}"
    output_path = TEMP_DIR / f"{output_id}.{target_format}"

    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)

        img = Image.open(input_path)

        if target_format in ('jpg', 'jpeg') and img.mode == 'RGBA':
            background = Image.new('RGB', img.size, (255, 255, 255))
            background.paste(img, mask=img.split()[3])
            img = background
        elif target_format in ('jpg', 'jpeg') and img.mode != 'RGB':
            img = img.convert('RGB')
        elif target_format == 'png' and img.mode not in ('RGBA', 'RGB'):
            img = img.convert('RGBA')

        if target_format in ('jpg', 'jpeg'):
            img.save(output_path, 'JPEG', quality=92)
        elif target_format == 'png':
            img.save(output_path, 'PNG')
        elif target_format == 'webp':
            img.save(output_path, 'WEBP', quality=92)

        cleanup_temp(input_path)

        media_type = 'image/jpeg' if target_format == 'jpg' else f'image/{target_format}'
        return FileResponse(
            path=output_path,
            media_type=media_type,
            filename=f"{os.path.splitext(file.filename)[0]}.{target_format}",
            headers={
                "X-Original-Size": str(len(content)),
                "X-Output-Size": str(os.path.getsize(output_path))
            }
        )
    except Exception as e:
        cleanup_temp(input_path, output_path)
        raise HTTPException(status_code=500, detail=f"Erro ao converter imagem: {str(e)}")


# ─── REDIMENSIONAR IMAGEM ─────────────────────────────────────────────────────

@app.post("/resize-image")
async def resize_image(
    file: UploadFile = File(...),
    width: int = Form(...),
    height: int = Form(...),
    keep_aspect: str = Form("true"),
):
    allowed_ext = {'.jpg', '.jpeg', '.png', '.webp'}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed_ext:
        raise HTTPException(status_code=400, detail="Formato nao suportado. Envie JPG, PNG ou WebP")

    if width < 1 or width > 10000 or height < 1 or height > 10000:
        raise HTTPException(status_code=400, detail="Dimensoes devem ser entre 1 e 10000 pixels")

    input_id = str(uuid.uuid4())
    output_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}{ext}"
    output_ext = ext if ext != '.jpeg' else '.jpg'
    output_path = TEMP_DIR / f"{output_id}_resized{output_ext}"

    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)

        img = Image.open(input_path)

        keep = keep_aspect.lower() in ('true', '1', 'yes')

        if keep:
            ratio = min(width / img.width, height / img.height)
            new_w = round(img.width * ratio)
            new_h = round(img.height * ratio)
        else:
            new_w = width
            new_h = height

        if new_w < 1:
            new_w = 1
        if new_h < 1:
            new_h = 1

        if new_w > 10000:
            new_w = 10000
        if new_h > 10000:
            new_h = 10000

        resized = img.resize((new_w, new_h), Image.LANCZOS)

        save_kwargs = {}
        if output_ext == '.jpg':
            save_kwargs = {'format': 'JPEG', 'quality': 92}
        elif output_ext == '.png':
            save_kwargs = {'format': 'PNG'}
        elif output_ext == '.webp':
            save_kwargs = {'format': 'WEBP', 'quality': 92}

        if resized.mode == 'RGBA' and output_ext == '.jpg':
            background = Image.new('RGB', resized.size, (255, 255, 255))
            background.paste(resized, mask=resized.split()[3])
            resized = background

        resized.save(output_path, **save_kwargs)

        cleanup_temp(input_path)

        return FileResponse(
            path=output_path,
            media_type=f"image/{output_ext.replace('.', '')}",
            filename=f"resized_{os.path.splitext(file.filename)[0]}{output_ext}",
            headers={
                "X-Original-Size": str(len(content)),
                "X-Output-Size": str(os.path.getsize(output_path))
            }
        )
    except Exception as e:
        cleanup_temp(input_path, output_path)
        raise HTTPException(status_code=500, detail=f"Erro ao redimensionar: {str(e)}")


# ─── CROP PDF ─────────────────────────────────────────────────────────────────

@app.post("/crop-pdf")
async def crop_pdf(
    file: UploadFile = File(...),
    x: float = Form(...),
    y: float = Form(...),
    width: float = Form(...),
    height: float = Form(...),
    pages: str = Form("all"),
):
    validate_pdf(file.filename)
    if width <= 0 or height <= 0:
        raise HTTPException(status_code=400, detail="Largura e altura devem ser maiores que zero")

    input_id = str(uuid.uuid4())
    output_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}.pdf"
    output_path = TEMP_DIR / f"{output_id}_cropped.pdf"

    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)

        import fitz
        doc = fitz.open(str(input_path))
        total_pages = doc.page_count

        page_list = parse_page_range(pages, total_pages)

        for i in range(total_pages):
            if (i + 1) in page_list:
                page = doc[i]
                page_rect = page.rect
                page_h = page_rect.height

                crop_top = y
                crop_left = x
                crop_bottom = y + height
                crop_right = x + width

                crop_left = max(0, min(crop_left, page_rect.width))
                crop_top = max(0, min(crop_top, page_rect.height))
                crop_right = max(crop_left + 1, min(crop_right, page_rect.width))
                crop_bottom = max(crop_top + 1, min(crop_bottom, page_rect.height))

                crop_rect = fitz.Rect(crop_left, crop_top, crop_right, crop_bottom)
                page.set_cropbox(crop_rect)

        doc.save(
            output_path,
            garbage=3,
            deflate=True,
        )
        doc.close()
        cleanup_temp(input_path)

        return FileResponse(
            path=output_path,
            media_type="application/pdf",
            filename=f"cropped_{file.filename}",
            headers={
                "X-Total-Pages": str(total_pages),
                "X-Cropped-Pages": str(len(page_list))
            }
        )
    except HTTPException:
        raise
    except Exception as e:
        cleanup_temp(input_path, output_path)
        raise HTTPException(status_code=500, detail=f"Erro ao cortar PDF: {str(e)}")


# ─── OCR PDF ──────────────────────────────────────────────────────────────────

@app.post("/ocr-pdf")
async def ocr_pdf(
    file: UploadFile = File(...),
    lang: str = Form("pt"),
):
    validate_pdf(file.filename)
    input_id = str(uuid.uuid4())
    output_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}.pdf"
    output_path = TEMP_DIR / f"{output_id}_ocr.txt"

    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)

        import fitz
        doc = fitz.open(str(input_path))
        total_pages = doc.page_count

        try:
            import paddleocr
            engine = paddleocr.PaddleOCR(lang=lang, use_angle_cls=True, show_log=False)
        except ImportError:
            try:
                from ocr.paddle.engine import get_engine
                engine = get_engine(lang=lang, use_gpu=False)
            except Exception:
                engine = None

        all_text = []

        for i in range(total_pages):
            page = doc[i]
            pix = page.get_pixmap(dpi=300)
            img_data = pix.tobytes("png")

            if engine:
                try:
                    if isinstance(engine, paddleocr.PaddleOCR):
                        result = engine.ocr(img_data, cls=True)
                    else:
                        import numpy as np
                        from PIL import Image as PILImage
                        img_array = np.array(PILImage.open(io.BytesIO(img_data)))
                        result = engine.ocr(img_array, cls=False)

                    page_text = []
                    if result and result[0]:
                        for line in result[0]:
                            if line and len(line) >= 2:
                                text = line[1][0] if isinstance(line[1], (list, tuple)) else line[1]
                                page_text.append(text)
                    all_text.append(f"--- Pagina {i + 1} ---\n" + "\n".join(page_text))
                except Exception as e:
                    all_text.append(f"--- Pagina {i + 1} ---\n[Erro no OCR: {str(e)}]")
            else:
                text = page.get_text()
                all_text.append(f"--- Pagina {i + 1} ---\n{text if text.strip() else '(Texto extraido nativamente)'}")

        doc.close()

        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n\n".join(all_text))

        cleanup_temp(input_path)

        return FileResponse(
            path=output_path,
            media_type="text/plain; charset=utf-8",
            filename=f"{file.filename.replace('.pdf', '')}_ocr.txt",
            headers={"X-Total-Pages": str(total_pages)}
        )
    except HTTPException:
        raise
    except Exception as e:
        cleanup_temp(input_path, output_path)
        raise HTTPException(status_code=500, detail=f"Erro no OCR: {str(e)}")


# ─── IMAGEM PARA TEXTO (OCR) ─────────────────────────────────────────────────

@app.post("/image-to-text")
async def image_to_text(
    file: UploadFile = File(...),
    lang: str = Form("pt"),
):
    allowed_ext = {'.jpg', '.jpeg', '.png', '.webp', '.bmp', '.tiff', '.tif'}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed_ext:
        raise HTTPException(status_code=400, detail="Formato nao suportado. Envie JPG, PNG, WebP ou BMP")

    input_id = str(uuid.uuid4())
    output_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}{ext}"
    output_path = TEMP_DIR / f"{output_id}_ocr.txt"

    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)

        try:
            from ocr.paddle.engine import recognize_text as paddle_recognize
            results = paddle_recognize(str(input_path), lang=lang, use_gpu=False)
            lines = [r['text'] for r in results if r.get('text')]
            text = "\n".join(lines)
        except Exception:
            try:
                import paddleocr
                engine = paddleocr.PaddleOCR(lang=lang, use_angle_cls=True, show_log=False)
                result = engine.ocr(str(input_path), cls=True)
                lines = []
                if result and result[0]:
                    for line in result[0]:
                        if line and len(line) >= 2:
                            t = line[1][0] if isinstance(line[1], (list, tuple)) else line[1]
                            lines.append(t)
                text = "\n".join(lines)
            except Exception:
                text = "[OCR nao disponivel. Instale PaddleOCR: pip install paddlepaddle paddleocr]"

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(text)

        cleanup_temp(input_path)

        return FileResponse(
            path=output_path,
            media_type="text/plain; charset=utf-8",
            filename=f"{os.path.splitext(file.filename)[0]}_ocr.txt"
        )
    except Exception as e:
        cleanup_temp(input_path, output_path)
        raise HTTPException(status_code=500, detail=f"Erro no OCR: {str(e)}")


# ─── MARCA D'AGUA ─────────────────────────────────────────────────────────────

@app.post("/watermark-pdf")
async def watermark_pdf(
    file: UploadFile = File(...),
    text: str = Form(""),
    opacity: float = Form(0.3),
    position: str = Form("center"),
    font_size: int = Form(48),
    pages: str = Form("all"),
    image: Optional[UploadFile] = File(None),
):
    validate_pdf(file.filename)
    if not text.strip() and image is None:
        raise HTTPException(status_code=400, detail="Informe um texto ou uma imagem para a marca dagua")

    input_id = str(uuid.uuid4())
    output_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}.pdf"
    output_path = TEMP_DIR / f"{output_id}_watermarked.pdf"

    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)

        import fitz
        doc = fitz.open(str(input_path))
        total_pages = doc.page_count
        page_list = parse_page_range(pages, total_pages)
        opacity_clamped = max(0.05, min(1.0, opacity))

        watermark_image = None
        if image:
            img_content = await image.read()
            watermark_image = io.BytesIO(img_content)

        for i in range(total_pages):
            if (i + 1) not in page_list:
                continue

            page = doc[i]
            page_rect = page.rect
            center_x = page_rect.width / 2
            center_y = page_rect.height / 2

            if watermark_image:
                if position == 'top':
                    img_y = page_rect.height * 0.05
                elif position == 'bottom':
                    img_y = page_rect.height * 0.7
                else:
                    img_y = center_y - 50

                watermark_image.seek(0)
                img_rect = fitz.Rect(
                    center_x - 60, img_y,
                    center_x + 60, img_y + 40
                )
                page.insert_image(img_rect, stream=watermark_image.read(), keep_proportion=True, alpha=int(opacity_clamped * 255))

            elif text.strip():
                if position == 'top':
                    pos_y = page_rect.height * 0.15
                elif position == 'bottom':
                    pos_y = page_rect.height * 0.85
                else:
                    pos_y = center_y

                rotation = 45 if position == 'diagonal' else 0
                text_width = fitz.get_text_length(text, fontname="helvetica", fontsize=font_size)

                rc = fitz.Rect(
                    center_x - text_width / 2 - 20, pos_y - font_size,
                    center_x + text_width / 2 + 20, pos_y + font_size
                )

                page.insert_textbox(
                    rc, text,
                    fontname="helvetica",
                    fontsize=font_size,
                    color=(0.5, 0.5, 0.5),
                    align=fitz.TEXT_ALIGN_CENTER,
                    rotate=rotation,
                    stroke_opacity=opacity_clamped,
                    fill_opacity=opacity_clamped,
                )

        doc.save(output_path, garbage=3, deflate=True)
        doc.close()
        cleanup_temp(input_path)

        return FileResponse(
            path=output_path,
            media_type="application/pdf",
            filename=f"watermarked_{file.filename}",
            headers={"X-Total-Pages": str(total_pages)}
        )
    except HTTPException:
        raise
    except Exception as e:
        cleanup_temp(input_path, output_path)
        raise HTTPException(status_code=500, detail=f"Erro ao aplicar marca dagua: {str(e)}")


# ─── NUMERAR PAGINAS ──────────────────────────────────────────────────────────

@app.post("/number-pages")
async def number_pages(
    file: UploadFile = File(...),
    start_number: int = Form(1),
    format_str: str = Form("1"),
    position: str = Form("bottom"),
    margin: int = Form(30),
    skip_first: str = Form("false"),
):
    validate_pdf(file.filename)
    input_id = str(uuid.uuid4())
    output_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}.pdf"
    output_path = TEMP_DIR / f"{output_id}_numbered.pdf"

    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)

        import fitz
        doc = fitz.open(str(input_path))
        total_pages = doc.page_count
        skip = skip_first.lower() in ('true', '1', 'yes')

        for i in range(total_pages):
            if skip and i == 0:
                continue

            page = doc[i]
            page_rect = page.rect
            num = start_number + (i - 1 if skip else i)
            page_total = total_pages - (1 if skip else 0)

            num_text = format_str
            num_text = num_text.replace('1', str(num))
            num_text = num_text.replace('{page}', str(num))
            num_text = num_text.replace('{total}', str(page_total))

            if '{' not in format_str:
                num_text = str(num) if format_str == '1' else format_str.replace('1', str(num), 1)

            font_size = 10
            text_width = fitz.get_text_length(num_text, fontname="helvetica", fontsize=font_size)

            if position == 'top':
                x = (page_rect.width - text_width) / 2
                y = margin
            elif position == 'top-right':
                x = page_rect.width - text_width - margin
                y = margin
            elif position == 'top-left':
                x = margin
                y = margin
            elif position == 'bottom-right':
                x = page_rect.width - text_width - margin
                y = page_rect.height - margin - font_size
            elif position == 'bottom-left':
                x = margin
                y = page_rect.height - margin - font_size
            else:
                x = (page_rect.width - text_width) / 2
                y = page_rect.height - margin - font_size

            page.insert_text(
                fitz.Point(x, y),
                num_text,
                fontname="helvetica",
                fontsize=font_size,
                color=(0, 0, 0)
            )

        doc.save(output_path)
        doc.close()
        cleanup_temp(input_path)

        return FileResponse(
            path=output_path,
            media_type="application/pdf",
            filename=f"numbered_{file.filename}",
            headers={"X-Total-Pages": str(total_pages)}
        )
    except HTTPException:
        raise
    except Exception as e:
        cleanup_temp(input_path, output_path)
        raise HTTPException(status_code=500, detail=f"Erro ao numerar paginas: {str(e)}")


# ─── PDF PARA WORD ────────────────────────────────────────────────────────────

# --- EDITAR PDF --------------------------------------------------------------

@app.post("/pdf-text-layer")
async def pdf_text_layer(file: UploadFile = File(...)):
    validate_pdf(file.filename)
    input_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}.pdf"

    doc = None
    try:
        content = await file.read()
        if len(content) > MAX_FILE_SIZE:
            raise HTTPException(status_code=413, detail="Arquivo excede 200MB")
        with open(input_path, "wb") as f:
            f.write(content)

        import fitz

        doc = fitz.open(str(input_path))
        pages = []
        total_items = 0

        for page_index in range(doc.page_count):
            page = doc[page_index]
            page_dict = page.get_text("dict")
            items = []

            for block in page_dict.get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line_index, line in enumerate(block.get("lines", [])):
                    spans = line.get("spans", [])
                    text = "".join(span.get("text", "") for span in spans).strip()
                    if not text:
                        continue

                    bbox = line.get("bbox") or spans[0].get("bbox")
                    if not bbox:
                        continue

                    sizes = [float(span.get("size", 12)) for span in spans if span.get("text", "").strip()]
                    first_span = next((span for span in spans if span.get("text", "").strip()), spans[0])
                    item = {
                        "id": f"p{page_index + 1}_b{block.get('number', 0)}_l{line_index}",
                        "page": page_index + 1,
                        "text": text,
                        "x": float(bbox[0]),
                        "y": float(bbox[1]),
                        "width": float(bbox[2] - bbox[0]),
                        "height": float(bbox[3] - bbox[1]),
                        "fontSize": round(sum(sizes) / len(sizes), 2) if sizes else 12,
                        "font": str(first_span.get("font", "helvetica")),
                        "color": pdf_color_to_hex(first_span.get("color")),
                    }
                    items.append(item)
                    total_items += 1

            pages.append({
                "page": page_index + 1,
                "width": float(page.rect.width),
                "height": float(page.rect.height),
                "items": items,
            })

        return {
            "filename": file.filename,
            "total_pages": doc.page_count,
            "text_items": total_items,
            "has_text_layer": total_items > 0,
            "pages": pages,
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao ler texto do PDF: {str(e)}")
    finally:
        if doc is not None:
            doc.close()
        cleanup_temp(input_path)


@app.post("/edit-pdf")
async def edit_pdf(
    file: UploadFile = File(...),
    operations: str = Form("[]"),
):
    validate_pdf(file.filename)

    try:
        parsed_operations = json.loads(operations)
    except json.JSONDecodeError:
        raise HTTPException(status_code=400, detail="Lista de operacoes invalida")

    if not isinstance(parsed_operations, list):
        raise HTTPException(status_code=400, detail="Operacoes devem ser um array JSON")
    if len(parsed_operations) > 500:
        raise HTTPException(status_code=400, detail="Limite de 500 operacoes por PDF")

    input_id = str(uuid.uuid4())
    output_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}.pdf"
    output_path = TEMP_DIR / f"{output_id}_edited.pdf"

    doc = None
    try:
        content = await file.read()
        if len(content) > MAX_FILE_SIZE:
            raise HTTPException(status_code=413, detail="Arquivo excede 200MB")
        with open(input_path, "wb") as f:
            f.write(content)

        import fitz

        doc = fitz.open(str(input_path))
        total_pages = doc.page_count
        used_field_names: set[str] = set()
        applied = 0
        fields_added = 0
        replaced_text = 0

        for index, raw_op in enumerate(parsed_operations):
            if not isinstance(raw_op, dict):
                raise HTTPException(status_code=400, detail=f"Operacao {index + 1} invalida")

            kind = str(raw_op.get("kind") or raw_op.get("type") or "").strip().lower()
            try:
                page_num = int(raw_op.get("page", 1))
            except (TypeError, ValueError):
                raise HTTPException(status_code=400, detail=f"Pagina invalida na operacao {index + 1}")

            if page_num < 1 or page_num > total_pages:
                raise HTTPException(status_code=400, detail=f"Pagina {page_num} fora do PDF")

            page = doc[page_num - 1]
            rect = operation_rect(fitz, raw_op, page.rect)

            if kind in ("text", "replace_text"):
                text = str(raw_op.get("text") or "").strip()
                if not text:
                    continue

                font_size = max(6, min(240, float(raw_op.get("fontSize", raw_op.get("font_size", 14)))))
                opacity = max(0.05, min(1.0, float(raw_op.get("opacity", 1))))
                align_map = {
                    "left": fitz.TEXT_ALIGN_LEFT,
                    "center": fitz.TEXT_ALIGN_CENTER,
                    "right": fitz.TEXT_ALIGN_RIGHT,
                }
                align = align_map.get(str(raw_op.get("align", "left")).lower(), fitz.TEXT_ALIGN_LEFT)
                color = parse_pdf_color(raw_op.get("color"), (0, 0, 0))
                background = parse_pdf_color(raw_op.get("backgroundColor"), (1, 1, 1))

                if kind == "replace_text":
                    pad = max(0, min(12, float(raw_op.get("padding", 1.5))))
                    redact_rect = fitz.Rect(
                        max(0, rect.x0 - pad),
                        max(0, rect.y0 - pad),
                        min(page.rect.width, rect.x1 + pad),
                        min(page.rect.height, rect.y1 + pad),
                    )
                    page.add_redact_annot(redact_rect, fill=background)
                    page.apply_redactions(images=0, graphics=0, text=0)
                    replaced_text += 1
                    if "\n" not in text:
                        baseline_y = max(rect.y0 + font_size, rect.y1 - max(1, font_size * 0.15))
                        page.insert_text(
                            fitz.Point(rect.x0, baseline_y),
                            text,
                            fontname="helvetica",
                            fontsize=font_size,
                            color=color,
                            fill_opacity=opacity,
                            stroke_opacity=opacity,
                        )
                        applied += 1
                        continue

                remaining = page.insert_textbox(
                    rect,
                    text,
                    fontname="helvetica",
                    fontsize=font_size,
                    color=color,
                    align=align,
                    fill_opacity=opacity,
                    stroke_opacity=opacity,
                )

                if remaining <= 0:
                    baseline_y = max(rect.y0 + font_size, rect.y1 - max(1, font_size * 0.15))
                    page.insert_text(
                        fitz.Point(rect.x0, baseline_y),
                        text,
                        fontname="helvetica",
                        fontsize=font_size,
                        color=color,
                        fill_opacity=opacity,
                        stroke_opacity=opacity,
                    )

                applied += 1
                continue

            widget_type_map = {
                "text_field": fitz.PDF_WIDGET_TYPE_TEXT,
                "field_text": fitz.PDF_WIDGET_TYPE_TEXT,
                "checkbox": fitz.PDF_WIDGET_TYPE_CHECKBOX,
                "radio": fitz.PDF_WIDGET_TYPE_RADIOBUTTON,
                "dropdown": fitz.PDF_WIDGET_TYPE_COMBOBOX,
                "combobox": fitz.PDF_WIDGET_TYPE_COMBOBOX,
                "listbox": fitz.PDF_WIDGET_TYPE_LISTBOX,
                "signature": fitz.PDF_WIDGET_TYPE_SIGNATURE,
            }

            if kind not in widget_type_map:
                raise HTTPException(status_code=400, detail=f"Tipo de operacao nao suportado: {kind}")

            prefix = {
                fitz.PDF_WIDGET_TYPE_TEXT: "campo_texto",
                fitz.PDF_WIDGET_TYPE_CHECKBOX: "checkbox",
                fitz.PDF_WIDGET_TYPE_RADIOBUTTON: "radio",
                fitz.PDF_WIDGET_TYPE_COMBOBOX: "lista",
                fitz.PDF_WIDGET_TYPE_LISTBOX: "lista",
                fitz.PDF_WIDGET_TYPE_SIGNATURE: "assinatura",
            }[widget_type_map[kind]]

            field_name_source = raw_op.get("groupName") if kind == "radio" else raw_op.get("name")
            if kind == "radio":
                raw_name = str(field_name_source or f"{prefix}_{index + 1}").strip()
                field_name = "".join(ch if ch.isalnum() or ch in ("_", "-", ".") else "_" for ch in raw_name)
                field_name = field_name[:80] or f"{prefix}_{index + 1}"
            else:
                field_name = safe_field_name(field_name_source, prefix, index, used_field_names)
            widget = fitz.Widget()
            widget.field_type = widget_type_map[kind]
            widget.field_name = field_name
            widget.field_label = str(raw_op.get("label") or raw_op.get("placeholder") or field_name)
            widget.rect = rect
            widget.border_color = parse_pdf_color(raw_op.get("borderColor"), (0.1, 0.1, 0.1))
            widget.fill_color = parse_pdf_color(raw_op.get("fillColor"), (1, 1, 1))
            widget.border_width = max(0, min(4, float(raw_op.get("borderWidth", 1))))
            widget.text_font = "Helv"
            widget.text_fontsize = max(0, min(72, float(raw_op.get("fontSize", raw_op.get("font_size", 11)))))
            widget.text_color = parse_pdf_color(raw_op.get("color"), (0, 0, 0))

            if parse_bool(raw_op.get("required", False)):
                widget.field_flags |= fitz.PDF_FIELD_IS_REQUIRED
            if parse_bool(raw_op.get("readOnly", False)):
                widget.field_flags |= fitz.PDF_FIELD_IS_READ_ONLY

            if widget.field_type == fitz.PDF_WIDGET_TYPE_TEXT:
                widget.field_value = str(raw_op.get("value") or "")
                if parse_bool(raw_op.get("multiline", False)):
                    widget.field_flags |= fitz.PDF_TX_FIELD_IS_MULTILINE
                max_len = raw_op.get("maxLength")
                if max_len not in (None, ""):
                    try:
                        widget.text_maxlen = max(0, int(max_len))
                    except (TypeError, ValueError):
                        pass

            elif widget.field_type == fitz.PDF_WIDGET_TYPE_CHECKBOX:
                widget.field_value = parse_bool(raw_op.get("checked", False))

            elif widget.field_type == fitz.PDF_WIDGET_TYPE_RADIOBUTTON:
                widget.field_label = str(raw_op.get("option") or raw_op.get("label") or field_name)
                widget.field_value = False

            elif widget.field_type in (fitz.PDF_WIDGET_TYPE_COMBOBOX, fitz.PDF_WIDGET_TYPE_LISTBOX):
                options = raw_op.get("options") or []
                if isinstance(options, str):
                    options = [item.strip() for item in options.splitlines() if item.strip()]
                if not isinstance(options, list) or not options:
                    options = ["Opcao 1", "Opcao 2", "Opcao 3"]
                options = [str(item)[:120] for item in options]
                widget.choice_values = options
                value = str(raw_op.get("value") or "")
                widget.field_value = value if value in options else options[0]

            elif widget.field_type == fitz.PDF_WIDGET_TYPE_SIGNATURE:
                widget.field_value = ""

            page.add_widget(widget)
            fields_added += 1
            applied += 1

        if fields_added:
            doc.need_appearances(True)

        doc.save(str(output_path), garbage=3, deflate=True)
        doc.close()
        doc = None
        cleanup_temp(input_path)

        return FileResponse(
            path=output_path,
            media_type="application/pdf",
            filename=f"edited_{file.filename}",
            headers={
                "X-Total-Pages": str(total_pages),
                "X-Operations-Applied": str(applied),
                "X-Fields-Added": str(fields_added),
                "X-Replaced-Text": str(replaced_text),
            }
        )
    except HTTPException:
        cleanup_temp(input_path, output_path)
        raise
    except Exception as e:
        cleanup_temp(input_path, output_path)
        raise HTTPException(status_code=500, detail=f"Erro ao editar PDF: {str(e)}")
    finally:
        if doc is not None:
            doc.close()
        cleanup_temp(input_path)


@app.post("/pdf-to-word")
async def pdf_to_word(
    file: UploadFile = File(...),
    pages: str = Form(""),
):
    validate_pdf(file.filename)
    input_id = str(uuid.uuid4())
    output_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}.pdf"
    output_path = TEMP_DIR / f"{output_id}.docx"

    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)

        reader = PdfReader(input_path)
        total_pages = len(reader.pages)
        page_list = parse_page_range(pages, total_pages) if pages.strip() else list(range(1, total_pages + 1))

        try:
            from pdf2docx import Converter as DocxConverter
            cv = DocxConverter(input_path)
            if len(page_list) == total_pages:
                cv.convert(str(output_path))
            else:
                cv.convert(str(output_path), pages=page_list)
            cv.close()
        except ImportError:
            import fitz
            doc = fitz.open(str(input_path))
            from docx import Document

            word_doc = Document()
            for page_num in page_list:
                page = doc[page_num - 1]
                text = page.get_text()
                if text.strip():
                    word_doc.add_paragraph(text)
                if page_num < page_list[-1]:
                    word_doc.add_page_break()
            doc.close()
            word_doc.save(str(output_path))

        cleanup_temp(input_path)

        return FileResponse(
            path=output_path,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{file.filename.replace('.pdf', '')}.docx",
            headers={"X-Pages-Processed": str(len(page_list))}
        )
    except HTTPException:
        raise
    except Exception as e:
        cleanup_temp(input_path, output_path)
        raise HTTPException(status_code=500, detail=f"Erro ao converter para Word: {str(e)}")


# ─── WORD PARA PDF ────────────────────────────────────────────────────────────

@app.post("/word-to-pdf")
async def word_to_pdf(file: UploadFile = File(...)):
    if not file.filename.lower().endswith(('.docx', '.doc')):
        raise HTTPException(status_code=400, detail="Envie um arquivo .docx ou .doc")

    input_id = str(uuid.uuid4())
    output_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}{os.path.splitext(file.filename)[1]}"
    output_path = TEMP_DIR / f"{output_id}.pdf"

    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)

        try:
            from docx2pdf import convert as docx2pdf_convert
            docx2pdf_convert(str(input_path), str(output_path))
        except ImportError:
            try:
                from docx import Document
                from reportlab.lib.pagesizes import A4
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                from reportlab.lib.styles import getSampleStyleSheet
                from reportlab.lib.units import mm

                doc = Document(str(input_path))
                pdf_doc = SimpleDocTemplate(str(output_path), pagesize=A4)
                styles = getSampleStyleSheet()
                elements = []

                for para in doc.paragraphs:
                    if para.text.strip():
                        elements.append(Paragraph(para.text, styles['Normal']))
                        elements.append(Spacer(1, 4 * mm))

                pdf_doc.build(elements)
            except Exception:
                cleanup_temp(input_path, output_path)
                raise HTTPException(status_code=500, detail="LibreOffice ou dependencias de conversao Word nao disponiveis. Instale: pip install docx2pdf ou use LibreOffice headless.")

        cleanup_temp(input_path)

        return FileResponse(
            path=output_path,
            media_type="application/pdf",
            filename=f"{os.path.splitext(file.filename)[0]}.pdf"
        )
    except HTTPException:
        raise
    except Exception as e:
        cleanup_temp(input_path, output_path)
        raise HTTPException(status_code=500, detail=f"Erro ao converter Word: {str(e)}")


# ─── METADADOS PDF ────────────────────────────────────────────────────────────

@app.post("/pdf-metadata")
async def pdf_metadata(file: UploadFile = File(...)):
    validate_pdf(file.filename)
    input_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}.pdf"

    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)

        reader = PdfReader(input_path)
        meta = reader.metadata or {}
        total_pages = len(reader.pages)

        result = {
            "title": str(meta.get('/Title', '')),
            "author": str(meta.get('/Author', '')),
            "subject": str(meta.get('/Subject', '')),
            "creator": str(meta.get('/Creator', '')),
            "producer": str(meta.get('/Producer', '')),
            "total_pages": total_pages,
            "file_size_bytes": len(content),
        }

        cleanup_temp(input_path)
        return result
    except Exception as e:
        cleanup_temp(input_path)
        raise HTTPException(status_code=500, detail=f"Erro ao ler metadados: {str(e)}")


@app.post("/edit-pdf-metadata")
async def edit_pdf_metadata(
    file: UploadFile = File(...),
    title: str = Form(""),
    author: str = Form(""),
    subject: str = Form(""),
):
    validate_pdf(file.filename)
    input_id = str(uuid.uuid4())
    output_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}.pdf"
    output_path = TEMP_DIR / f"{output_id}_meta.pdf"

    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)

        reader = PdfReader(input_path)
        writer = PdfWriter()

        for page in reader.pages:
            writer.add_page(page)

        metadata = {}
        if title:
            metadata['/Title'] = title
        if author:
            metadata['/Author'] = author
        if subject:
            metadata['/Subject'] = subject

        writer.add_metadata(metadata)

        with open(output_path, "wb") as of:
            writer.write(of)

        cleanup_temp(input_path)

        return FileResponse(
            path=output_path,
            media_type="application/pdf",
            filename=f"meta_{file.filename}"
        )
    except Exception as e:
        cleanup_temp(input_path, output_path)
        raise HTTPException(status_code=500, detail=f"Erro ao editar metadados: {str(e)}")


# ─── XML PARA EXCEL ───────────────────────────────────────────────────────────

@app.post("/xml-to-excel")
async def xml_to_excel(files: List[UploadFile] = File(...)):
    if not files:
        raise HTTPException(status_code=400, detail="Envie pelo menos um arquivo XML")

    import xml.etree.ElementTree as ET

    output_id = str(uuid.uuid4())
    output_path = TEMP_DIR / f"{output_id}.xlsx"

    try:
        import pandas as pd

        all_notas = []
        all_itens = []

        for f in files:
            content = await f.read()
            try:
                root = ET.fromstring(content)
            except ET.ParseError as e:
                raise HTTPException(status_code=400, detail=f"XML invalido: {f.filename} - {str(e)}")

            ns = {'nfe': 'http://www.portalcatalao.net/nfe'}

            def find_text(element, *paths):
                for path in paths:
                    found = element.find(path)
                    if found is not None and found.text:
                        return found.text
                return ''

            chave = find_text(root, './/{http://www.portalcatalao.net/nfe}chNFe', './/chNFe', './/infNFe', 'chNFe')
            numero = find_text(root, './/{http://www.portalcatalao.net/nfe}nNF', './/nNF')
            serie = find_text(root, './/{http://www.portalcatalao.net/nfe}serie', './/serie')
            dh_emi = find_text(root, './/{http://www.portalcatalao.net/nfe}dhEmi', './/dhEmi', './/dEmi')
            emitente = find_text(root, './/{http://www.portalcatalao.net/nfe}xNome', './/emit', './/xNome')
            destinatario_elem = root.find('.//{http://www.portalcatalao.net/nfe}dest')
            if destinatario_elem is not None:
                destinatario = find_text(destinatario_elem, './/{http://www.portalcatalao.net/nfe}xNome', './/xNome')
            else:
                destinatario = ''
            vnf = find_text(root, './/{http://www.portalcatalao.net/nfe}vNF', './/vNF')

            nota = {
                'chave': chave, 'numero': numero, 'serie': serie,
                'data_emissao': dh_emi, 'emitente': emitente,
                'destinatario': destinatario, 'valor_total': vnf
            }
            all_notas.append(nota)

            for item in root.findall('.//{http://www.portalcatalao.net/nfe}det'):
                cprod = find_text(item, './/{http://www.portalcatalao.net/nfe}cProd', './/cProd')
                xprod = find_text(item, './/{http://www.portalcatalao.net/nfe}xProd', './/xProd')
                qcom = find_text(item, './/{http://www.portalcatalao.net/nfe}qCom', './/qCom')
                vuncom = find_text(item, './/{http://www.portalcatalao.net/nfe}vUnCom', './/vUnCom')
                all_itens.append({
                    'chave': chave, 'numero': numero,
                    'codigo': cprod, 'descricao': xprod,
                    'quantidade': qcom, 'valor_unitario': vuncom
                })

        with pd.ExcelWriter(output_path, engine='openpyxl') as writer_obj:
            if all_notas:
                pd.DataFrame(all_notas).to_excel(writer_obj, sheet_name='notas', index=False)
            if all_itens:
                pd.DataFrame(all_itens).to_excel(writer_obj, sheet_name='itens', index=False)

        return FileResponse(
            path=output_path,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            filename="notas_fiscais.xlsx",
            headers={"X-Files-Processed": str(len(files))}
        )
    except HTTPException:
        raise
    except Exception as e:
        cleanup_temp(output_path)
        raise HTTPException(status_code=500, detail=f"Erro ao processar XML: {str(e)}")


# ─── VISUALIZAR XML ───────────────────────────────────────────────────────────

@app.post("/xml-preview")
async def xml_preview(file: UploadFile = File(...)):
    if not file.filename.lower().endswith('.xml'):
        raise HTTPException(status_code=400, detail="Envie um arquivo XML")

    import xml.etree.ElementTree as ET
    content = await file.read()

    try:
        root = ET.fromstring(content)
    except ET.ParseError as e:
        raise HTTPException(status_code=400, detail=f"XML invalido: {str(e)}")

    result = {
        "filename": file.filename,
        "valid": True,
    }

    def extract_nfe_data(xml_root):
        ns = 'http://www.portalcatalao.net/nfe'
        inf_nfe = xml_root.find(f'.//{{{ns}}}infNFe')
        if inf_nfe is None:
            return None

        def t(path):
            el = inf_nfe.find(f'.//{{{ns}}}{path}')
            return el.text if el is not None and el.text else ''

        return {
            'chave': inf_nfe.get('Id', '').replace('NFe', ''),
            'numero': t('ide/nNF'),
            'serie': t('ide/serie'),
            'data_emissao': t('ide/dhEmi') or t('ide/dEmi'),
            'emitente': {
                'cnpj': t('emit/CNPJ'),
                'nome': t('emit/xNome'),
                'fantasia': t('emit/xFant'),
            },
            'destinatario': {
                'cnpj': t('dest/CNPJ') or t('dest/CPF'),
                'nome': t('dest/xNome'),
            },
            'itens': [],
            'totais': {
                'vBC': t('total/ICMSTot/vBC'),
                'vICMS': t('total/ICMSTot/vICMS'),
                'vNF': t('total/ICMSTot/vNF'),
            }
        }

    nfe_data = extract_nfe_data(root)
    if nfe_data:
        result['tipo'] = 'NF-e'
        result['dados'] = nfe_data
    else:
        result['tipo'] = 'XML generico'
        result['mensagem'] = 'XML valido, mas nao identificado como NF-e'

    return result


# ─── VALIDAR XML ──────────────────────────────────────────────────────────────

@app.post("/validate-xml")
async def validate_xml(file: UploadFile = File(...)):
    if not file.filename.lower().endswith('.xml'):
        raise HTTPException(status_code=400, detail="Envie um arquivo XML")

    import xml.etree.ElementTree as ET
    content = await file.read()

    result = {
        "filename": file.filename,
        "file_size_bytes": len(content),
    }

    try:
        ET.fromstring(content)
        result["well_formed"] = True
        result["message"] = "XML bem formado e valido"
    except ET.ParseError as e:
        result["well_formed"] = False
        result["message"] = f"XML invalido: {str(e)}"
        result["error_line"] = getattr(e, 'position', None)

    return result


# ─── VISUALIZAR INFORMAÇÕES DO PDF ────────────────────────────────────────────

@app.post("/pdf-info")
async def pdf_info(file: UploadFile = File(...)):
    validate_pdf(file.filename)
    input_id = str(uuid.uuid4())
    input_path = TEMP_DIR / f"{input_id}.pdf"

    try:
        content = await file.read()
        with open(input_path, "wb") as f:
            f.write(content)

        reader = PdfReader(input_path)
        total_pages = len(reader.pages)
        meta = reader.metadata or {}
        encrypted = reader.is_encrypted

        cleanup_temp(input_path)

        return {
            "filename": file.filename,
            "total_pages": total_pages,
            "file_size_bytes": len(content),
            "encrypted": encrypted,
            "title": str(meta.get('/Title', '')),
            "author": str(meta.get('/Author', '')),
            "subject": str(meta.get('/Subject', '')),
            "creator": str(meta.get('/Creator', '')),
            "producer": str(meta.get('/Producer', '')),
        }
    except Exception as e:
        cleanup_temp(input_path)
        raise HTTPException(status_code=500, detail=f"Erro ao ler informacoes: {str(e)}")


# ─── CLEANUP ──────────────────────────────────────────────────────────────────

import threading
import time


def periodic_cleanup(interval_seconds: int = 3600):
    while True:
        time.sleep(interval_seconds)
        try:
            for f in TEMP_DIR.glob("*"):
                try:
                    if f.is_file():
                        f.unlink()
                    elif f.is_dir():
                        shutil.rmtree(f)
                except Exception:
                    pass
        except Exception:
            pass


cleanup_thread = threading.Thread(target=periodic_cleanup, daemon=True)
cleanup_thread.start()


@app.on_event("shutdown")
async def cleanup():
    for f in TEMP_DIR.glob("*"):
        try:
            if f.is_file():
                f.unlink()
            elif f.is_dir():
                shutil.rmtree(f)
        except Exception:
            pass
