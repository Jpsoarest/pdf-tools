import io
import json
from pathlib import Path
from typing import Any, List, Set, Tuple

from pypdf import PdfReader, PdfWriter

from core.tempfiles import cleanup_temp
from core.utils import (
    parse_page_range, parse_pdf_color, parse_bool,
    safe_field_name, operation_rect, pdf_color_to_hex
)
from core.errors import raise_bad_request, raise_internal_error


def get_text_layer_service(input_path: Path, filename: str) -> dict:
    try:
        import fitz
    except ImportError:
        raise_internal_error("Dependencia nao instalada: pymupdf")

    doc = fitz.open(str(input_path))
    pages = []
    total_items = 0

    for page_index in range(doc.page_count):
        page = doc[page_index]
        page_dict = page.get_text("dict")
        items = []

        for block in page_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line_index, line in enumerate(block.get("lines", [])):
                spans = line.get("spans", [])
                text = "".join(span.get("text", "") for span in spans).strip()
                if not text:
                    continue

                bbox = line.get("bbox") or spans[0].get("bbox")
                if not bbox:
                    continue

                sizes = [float(span.get("size", 12)) for span in spans if span.get("text", "").strip()]
                first_span = next((span for span in spans if span.get("text", "").strip()), spans[0])
                item = {
                    "id": f"p{page_index + 1}_b{block.get('number', 0)}_l{line_index}",
                    "page": page_index + 1,
                    "text": text,
                    "x": float(bbox[0]),
                    "y": float(bbox[1]),
                    "width": float(bbox[2] - bbox[0]),
                    "height": float(bbox[3] - bbox[1]),
                    "fontSize": round(sum(sizes) / len(sizes), 2) if sizes else 12,
                    "font": str(first_span.get("font", "helvetica")),
                    "color": pdf_color_to_hex(first_span.get("color")),
                }
                items.append(item)
                total_items += 1

        pages.append({
            "page": page_index + 1,
            "width": float(page.rect.width),
            "height": float(page.rect.height),
            "items": items,
        })

    doc.close()
    return {
        "filename": filename,
        "total_pages": doc.page_count if hasattr(doc, 'page_count') else len(pages),
        "text_items": total_items,
        "has_text_layer": total_items > 0,
        "pages": pages,
    }


def edit_pdf_service(
    input_path: Path, output_path: Path,
    operations: List[dict], filename: str
) -> Tuple[int, int, int]:
    try:
        import fitz
    except ImportError:
        raise_internal_error("Dependencia nao instalada: pymupdf")

    doc = fitz.open(str(input_path))
    total_pages = doc.page_count
    used_field_names: Set[str] = set()
    applied = 0
    fields_added = 0
    replaced_text = 0

    for index, raw_op in enumerate(operations):
        if not isinstance(raw_op, dict):
            raise_bad_request(f"Operacao {index + 1} invalida")

        kind = str(raw_op.get("kind") or raw_op.get("type") or "").strip().lower()
        try:
            page_num = int(raw_op.get("page", 1))
        except (TypeError, ValueError):
            raise_bad_request(f"Pagina invalida na operacao {index + 1}")

        if page_num < 1 or page_num > total_pages:
            raise_bad_request(f"Pagina {page_num} fora do PDF")

        page = doc[page_num - 1]
        rect = operation_rect(fitz, raw_op, page.rect)

        if kind in ("text", "replace_text"):
            text = str(raw_op.get("text") or "").strip()
            if not text:
                continue

            font_size = max(6, min(240, float(raw_op.get("fontSize", raw_op.get("font_size", 14)))))
            opacity = max(0.05, min(1.0, float(raw_op.get("opacity", 1))))
            align_map = {
                "left": fitz.TEXT_ALIGN_LEFT,
                "center": fitz.TEXT_ALIGN_CENTER,
                "right": fitz.TEXT_ALIGN_RIGHT,
            }
            align = align_map.get(str(raw_op.get("align", "left")).lower(), fitz.TEXT_ALIGN_LEFT)
            color = parse_pdf_color(raw_op.get("color"), (0, 0, 0))
            background = parse_pdf_color(raw_op.get("backgroundColor"), (1, 1, 1))

            if kind == "replace_text":
                pad = max(0, min(12, float(raw_op.get("padding", 1.5))))
                redact_rect = fitz.Rect(
                    max(0, rect.x0 - pad),
                    max(0, rect.y0 - pad),
                    min(page.rect.width, rect.x1 + pad),
                    min(page.rect.height, rect.y1 + pad),
                )
                page.add_redact_annot(redact_rect, fill=background)
                page.apply_redactions(images=0, graphics=0, text=0)
                replaced_text += 1
                if "\n" not in text:
                    baseline_y = max(rect.y0 + font_size, rect.y1 - max(1, font_size * 0.15))
                    page.insert_text(
                        fitz.Point(rect.x0, baseline_y),
                        text,
                        fontname="helvetica",
                        fontsize=font_size,
                        color=color,
                        fill_opacity=opacity,
                        stroke_opacity=opacity,
                    )
                    applied += 1
                    continue

            remaining = page.insert_textbox(
                rect,
                text,
                fontname="helvetica",
                fontsize=font_size,
                color=color,
                align=align,
                fill_opacity=opacity,
                stroke_opacity=opacity,
            )

            if remaining <= 0:
                baseline_y = max(rect.y0 + font_size, rect.y1 - max(1, font_size * 0.15))
                page.insert_text(
                    fitz.Point(rect.x0, baseline_y),
                    text,
                    fontname="helvetica",
                    fontsize=font_size,
                    color=color,
                    fill_opacity=opacity,
                    stroke_opacity=opacity,
                )

            applied += 1
            continue

        widget_type_map = {
            "text_field": fitz.PDF_WIDGET_TYPE_TEXT,
            "field_text": fitz.PDF_WIDGET_TYPE_TEXT,
            "checkbox": fitz.PDF_WIDGET_TYPE_CHECKBOX,
            "radio": fitz.PDF_WIDGET_TYPE_RADIOBUTTON,
            "dropdown": fitz.PDF_WIDGET_TYPE_COMBOBOX,
            "combobox": fitz.PDF_WIDGET_TYPE_COMBOBOX,
            "listbox": fitz.PDF_WIDGET_TYPE_LISTBOX,
            "signature": fitz.PDF_WIDGET_TYPE_SIGNATURE,
        }

        if kind not in widget_type_map:
            raise_bad_request(f"Tipo de operacao nao suportado: {kind}")

        prefix_map = {
            fitz.PDF_WIDGET_TYPE_TEXT: "campo_texto",
            fitz.PDF_WIDGET_TYPE_CHECKBOX: "checkbox",
            fitz.PDF_WIDGET_TYPE_RADIOBUTTON: "radio",
            fitz.PDF_WIDGET_TYPE_COMBOBOX: "lista",
            fitz.PDF_WIDGET_TYPE_LISTBOX: "lista",
            fitz.PDF_WIDGET_TYPE_SIGNATURE: "assinatura",
        }
        prefix = prefix_map[widget_type_map[kind]]

        field_name_source = raw_op.get("groupName") if kind == "radio" else raw_op.get("name")
        if kind == "radio":
            raw_name = str(field_name_source or f"{prefix}_{index + 1}").strip()
            field_name = "".join(ch if ch.isalnum() or ch in ("_", "-", ".") else "_" for ch in raw_name)
            field_name = field_name[:80] or f"{prefix}_{index + 1}"
        else:
            field_name = safe_field_name(field_name_source, prefix, index, used_field_names)

        widget = fitz.Widget()
        widget.field_type = widget_type_map[kind]
        widget.field_name = field_name
        widget.field_label = str(raw_op.get("label") or raw_op.get("placeholder") or field_name)
        widget.rect = rect
        widget.border_color = parse_pdf_color(raw_op.get("borderColor"), (0.1, 0.1, 0.1))
        widget.fill_color = parse_pdf_color(raw_op.get("fillColor"), (1, 1, 1))
        widget.border_width = max(0, min(4, float(raw_op.get("borderWidth", 1))))
        widget.text_font = "Helv"
        widget.text_fontsize = max(0, min(72, float(raw_op.get("fontSize", raw_op.get("font_size", 11)))))
        widget.text_color = parse_pdf_color(raw_op.get("color"), (0, 0, 0))

        if parse_bool(raw_op.get("required", False)):
            widget.field_flags |= fitz.PDF_FIELD_IS_REQUIRED
        if parse_bool(raw_op.get("readOnly", False)):
            widget.field_flags |= fitz.PDF_FIELD_IS_READ_ONLY

        if widget.field_type == fitz.PDF_WIDGET_TYPE_TEXT:
            widget.field_value = str(raw_op.get("value") or "")
            if parse_bool(raw_op.get("multiline", False)):
                widget.field_flags |= fitz.PDF_TX_FIELD_IS_MULTILINE
            max_len = raw_op.get("maxLength")
            if max_len not in (None, ""):
                try:
                    widget.text_maxlen = max(0, int(max_len))
                except (TypeError, ValueError):
                    pass

        elif widget.field_type == fitz.PDF_WIDGET_TYPE_CHECKBOX:
            widget.field_value = parse_bool(raw_op.get("checked", False))

        elif widget.field_type == fitz.PDF_WIDGET_TYPE_RADIOBUTTON:
            widget.field_label = str(raw_op.get("option") or raw_op.get("label") or field_name)
            widget.field_value = False

        elif widget.field_type in (fitz.PDF_WIDGET_TYPE_COMBOBOX, fitz.PDF_WIDGET_TYPE_LISTBOX):
            options = raw_op.get("options") or []
            if isinstance(options, str):
                options = [item.strip() for item in options.splitlines() if item.strip()]
            if not isinstance(options, list) or not options:
                options = ["Opcao 1", "Opcao 2", "Opcao 3"]
            options = [str(item)[:120] for item in options]
            widget.choice_values = options
            value = str(raw_op.get("value") or "")
            widget.field_value = value if value in options else options[0]

        elif widget.field_type == fitz.PDF_WIDGET_TYPE_SIGNATURE:
            widget.field_value = ""

        page.add_widget(widget)
        fields_added += 1
        applied += 1

    if fields_added:
        doc.need_appearances(True)

    doc.save(str(output_path), garbage=3, deflate=True)
    doc.close()
    cleanup_temp(input_path)
    return total_pages, applied, fields_added, replaced_text


def watermark_pdf_service(
    input_path: Path, output_path: Path,
    text: str, opacity: float, position: str,
    font_size: int, pages: str, image_data: bytes = None
) -> int:
    try:
        import fitz
    except ImportError:
        raise_internal_error("Dependencia nao instalada: pymupdf")

    doc = fitz.open(str(input_path))
    total_pages = doc.page_count
    page_list = parse_page_range(pages, total_pages)
    opacity_clamped = max(0.05, min(1.0, opacity))

    watermark_image = None
    if image_data:
        watermark_image = io.BytesIO(image_data)

    for i in range(total_pages):
        if (i + 1) not in page_list:
            continue

        page = doc[i]
        page_rect = page.rect
        center_x = page_rect.width / 2
        center_y = page_rect.height / 2

        if watermark_image:
            if position == 'top':
                img_y = page_rect.height * 0.05
            elif position == 'bottom':
                img_y = page_rect.height * 0.7
            else:
                img_y = center_y - 50

            watermark_image.seek(0)
            img_rect = fitz.Rect(center_x - 60, img_y, center_x + 60, img_y + 40)
            page.insert_image(img_rect, stream=watermark_image.read(), keep_proportion=True, alpha=int(opacity_clamped * 255))

        elif text.strip():
            if position == 'top':
                pos_y = page_rect.height * 0.15
            elif position == 'bottom':
                pos_y = page_rect.height * 0.85
            else:
                pos_y = center_y

            rotation = 45 if position == 'diagonal' else 0
            text_width = fitz.get_text_length(text, fontname="helvetica", fontsize=font_size)

            rc = fitz.Rect(
                center_x - text_width / 2 - 20, pos_y - font_size,
                center_x + text_width / 2 + 20, pos_y + font_size
            )

            page.insert_textbox(
                rc, text,
                fontname="helvetica",
                fontsize=font_size,
                color=(0.5, 0.5, 0.5),
                align=fitz.TEXT_ALIGN_CENTER,
                rotate=rotation,
                stroke_opacity=opacity_clamped,
                fill_opacity=opacity_clamped,
            )

    doc.save(output_path, garbage=3, deflate=True)
    doc.close()
    cleanup_temp(input_path)
    return total_pages


def number_pages_service(
    input_path: Path, output_path: Path,
    start_number: int, format_str: str,
    position: str, margin: int, skip_first: bool
) -> int:
    try:
        import fitz
    except ImportError:
        raise_internal_error("Dependencia nao instalada: pymupdf")

    doc = fitz.open(str(input_path))
    total_pages = doc.page_count

    for i in range(total_pages):
        if skip_first and i == 0:
            continue

        page = doc[i]
        page_rect = page.rect
        num = start_number + (i - 1 if skip_first else i)
        page_total = total_pages - (1 if skip_first else 0)

        if '{' not in format_str:
            num_text = str(num) if format_str == '1' else format_str.replace('1', str(num), 1)
        else:
            num_text = format_str.replace('{page}', str(num)).replace('{total}', str(page_total))

        font_size = 10
        text_width = fitz.get_text_length(num_text, fontname="helvetica", fontsize=font_size)

        if position == 'top':
            x = (page_rect.width - text_width) / 2
            y = margin
        elif position == 'top-right':
            x = page_rect.width - text_width - margin
            y = margin
        elif position == 'top-left':
            x = margin
            y = margin
        elif position == 'bottom-right':
            x = page_rect.width - text_width - margin
            y = page_rect.height - margin - font_size
        elif position == 'bottom-left':
            x = margin
            y = page_rect.height - margin - font_size
        else:
            x = (page_rect.width - text_width) / 2
            y = page_rect.height - margin - font_size

        page.insert_text(
            fitz.Point(x, y),
            num_text,
            fontname="helvetica",
            fontsize=font_size,
            color=(0, 0, 0)
        )

    doc.save(output_path)
    doc.close()
    cleanup_temp(input_path)
    return total_pages


def ocr_pdf_service(input_path: Path, output_path: Path, lang: str) -> int:
    try:
        import fitz
    except ImportError:
        raise_internal_error("Dependencia nao instalada: pymupdf")

    doc = fitz.open(str(input_path))
    total_pages = doc.page_count

    engine = None
    try:
        import paddleocr
        engine = paddleocr.PaddleOCR(lang=lang, use_angle_cls=True, show_log=False)
    except ImportError:
        try:
            from ocr.paddle.engine import get_engine
            engine = get_engine(lang=lang, use_gpu=False)
        except Exception:
            pass

    all_text = []

    for i in range(total_pages):
        page = doc[i]
        pix = page.get_pixmap(dpi=300)
        img_data = pix.tobytes("png")

        if engine:
            try:
                if 'paddleocr' in str(type(engine)).lower():
                    result = engine.ocr(img_data, cls=True)
                else:
                    import numpy as np
                    from PIL import Image as PILImage
                    img_array = np.array(PILImage.open(io.BytesIO(img_data)))
                    result = engine.ocr(img_array, cls=False)

                page_text = []
                if result and result[0]:
                    for line in result[0]:
                        if line and len(line) >= 2:
                            text = line[1][0] if isinstance(line[1], (list, tuple)) else line[1]
                            page_text.append(text)
                all_text.append(f"--- Pagina {i + 1} ---\n" + "\n".join(page_text))
            except Exception as e:
                all_text.append(f"--- Pagina {i + 1} ---\n[Erro no OCR: {str(e)}]")
        else:
            text = page.get_text()
            all_text.append(f"--- Pagina {i + 1} ---\n{text if text.strip() else '(Texto extraido nativamente)'}")

    doc.close()

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n\n".join(all_text))

    cleanup_temp(input_path)
    return total_pages


def image_to_text_service(input_path: Path, output_path: Path, lang: str) -> None:
    try:
        from ocr.paddle.engine import recognize_text as paddle_recognize
        results = paddle_recognize(str(input_path), lang=lang, use_gpu=False)
        lines = [r['text'] for r in results if r.get('text')]
        text = "\n".join(lines)
    except Exception:
        try:
            import paddleocr
            engine = paddleocr.PaddleOCR(lang=lang, use_angle_cls=True, show_log=False)
            result = engine.ocr(str(input_path), cls=True)
            lines = []
            if result and result[0]:
                for line in result[0]:
                    if line and len(line) >= 2:
                        t = line[1][0] if isinstance(line[1], (list, tuple)) else line[1]
                        lines.append(t)
            text = "\n".join(lines)
        except Exception:
            text = "[OCR nao disponivel. Instale PaddleOCR: pip install paddlepaddle paddleocr]"

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(text)

    cleanup_temp(input_path)


def pdf_to_word_service(input_path: Path, output_path: Path, pages: str) -> int:
    reader = PdfReader(input_path)
    total_pages = len(reader.pages)
    page_list = parse_page_range(pages, total_pages) if pages.strip() else list(range(1, total_pages + 1))

    try:
        from pdf2docx import Converter as DocxConverter
        cv = DocxConverter(input_path)
        if len(page_list) == total_pages:
            cv.convert(str(output_path))
        else:
            cv.convert(str(output_path), pages=page_list)
        cv.close()
    except ImportError:
        try:
            import fitz
            from docx import Document

            doc = fitz.open(str(input_path))
            word_doc = Document()
            for page_num in page_list:
                page = doc[page_num - 1]
                text = page.get_text()
                if text.strip():
                    word_doc.add_paragraph(text)
                if page_num < page_list[-1]:
                    word_doc.add_page_break()
            doc.close()
            word_doc.save(str(output_path))
        except ImportError:
            raise_internal_error("Dependencias nao instaladas: pdf2docx ou python-docx")

    cleanup_temp(input_path)
    return len(page_list)


def word_to_pdf_service(input_path: Path, output_path: Path) -> None:
    try:
        from docx2pdf import convert as docx2pdf_convert
        docx2pdf_convert(str(input_path), str(output_path))
    except ImportError:
        try:
            from docx import Document
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import mm

            doc = Document(str(input_path))
            pdf_doc = SimpleDocTemplate(str(output_path), pagesize=A4)
            styles = getSampleStyleSheet()
            elements = []

            for para in doc.paragraphs:
                if para.text.strip():
                    elements.append(Paragraph(para.text, styles['Normal']))
                    elements.append(Spacer(1, 4 * mm))

            pdf_doc.build(elements)
        except Exception:
            raise_internal_error(
                "LibreOffice ou dependencias de conversao Word nao disponiveis. "
                "Instale: pip install docx2pdf ou use LibreOffice headless."
            )

    cleanup_temp(input_path)
